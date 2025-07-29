import os
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from datetime import datetime
import logging
from django.conf import settings
from api.models import QuestionAnswer

# Configure logging
logging.basicConfig(level=logging.INFO)

class QuestionPaperGenerator:
    @staticmethod
    def format_table_cell(cell, text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
        paragraph = cell.paragraphs[0]
        paragraph.alignment = alignment
        run = paragraph.add_run(text)
        return run

    @staticmethod
    def add_equation_to_docx(paragraph, equation_data):
        try:
            math_element = parse_xml(equation_data['mathml'])
            paragraph._p.append(math_element)
            run = paragraph.add_run()
            run.add_text(" ")
        except Exception as e:
            logging.error(f"Error adding equation: {e}")
            paragraph.add_run(equation_data.get('text', 'Equation Error'))

    @staticmethod
    def add_image_to_docx(doc, cell, image_path):
        try:
            paragraph = cell.paragraphs[0]
            if os.path.exists(image_path):
                run = paragraph.add_run()
                picture = run.add_picture(image_path)
                
                # Fixed dimensions for all images
                fixed_width = Inches(3.0)  # 3 inches width
                fixed_height = Inches(2.0)  # 2 inches height
                
                picture.width = fixed_width
                picture.height = fixed_height
                
                # Add some spacing after the image
                run.add_text('\n')
        except Exception as e:
            logging.error(f"Error adding image: {e}")
            paragraph.add_run(f"[Image: {os.path.basename(image_path)}]")

    @staticmethod
    def create_paper(metadata, selected_questions, questions_data):
        try:
            # Use the exact template
            template_path = os.path.join(settings.BASE_DIR, 'template', 'generate_paper_template.docx')
            if not os.path.exists(template_path):
                logging.error(f"Template file not found at: {template_path}")
                raise FileNotFoundError(f"Template file not found at: {template_path}")
                
            doc = Document(template_path)
            
            # Clear any existing content after the header and academic year
            while len(doc.paragraphs) > 2:  # Keep header and academic year
                p = doc.paragraphs[-1]._element
                p.getparent().remove(p)
                p._p = p._element = None
            
            # Add department heading in red
            dept_heading = doc.add_paragraph()
            # Get department based on course title
            course_words = metadata.course_title.split()
            if any(word.lower() in ['computer', 'information'] for word in course_words):
                department = "INFORMATION SCIENCE AND ENGINEERING"
            elif any(word.lower() in ['ai', 'artificial', 'ml', 'machine'] for word in course_words):
                department = "ARTIFICIAL INTELLIGENCE AND MACHINE LEARNING"
            elif any(word.lower() in ['data', 'analytics'] for word in course_words):
                department = "DATA SCIENCE"
            else:
                department = "COMPUTER SCIENCE AND ENGINEERING"  # default department
                
            dept_run = dept_heading.add_run(f'DEPARTMENT OF {department}')
            dept_run.font.color.rgb = RGBColor(179, 0, 0)  # Dark red color
            dept_run.bold = True
            dept_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Create header table with specific formatting
            table = doc.add_table(rows=5, cols=4)
            table.style = 'Table Grid'
            table.allow_autofit = False  # Changed from autofit to allow_autofit
            
            # Set column widths
            for row in table.rows:
                row.cells[0].width = Inches(1.5)
                row.cells[1].width = Inches(2.5)
                row.cells[2].width = Inches(1.5)
                row.cells[3].width = Inches(2.5)

            # Fill header information
            header_data = [
                ['Date', metadata.date.strftime('%d-%m-%Y') if metadata.date else '', 'Maximum Marks', str(metadata.max_marks)],
                ['Course Code', metadata.course_code, 'Duration', metadata.duration],
                ['Sem', metadata.semester, 'Type', metadata.exam_type if hasattr(metadata, 'exam_type') else 'CIE'],  # Default to 'CIE' if not specified
                ['UG/PG', 'UG', 'Faculty:', metadata.faculty.name if metadata.faculty else ''],
            ]

            # Add the first 4 rows
            for i, row_data in enumerate(header_data):
                for j, text in enumerate(row_data):
                    cell = table.cell(i, j)
                    run = QuestionPaperGenerator.format_table_cell(cell, text)
                    if j == 0:  # Labels in first column
                        run.bold = True

            # Handle Course Title row separately
            course_title_cell = table.cell(4, 0)
            course_title_cell.merge(table.cell(4, 3))  # Merge all cells in last row
            
            # Remove borders from merged cell
            tc_pr = course_title_cell._tc.get_or_add_tcPr()
            tc_borders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>')
            tc_pr.append(tc_borders)

            # Add Course Title in red
            title_para = course_title_cell.paragraphs[0]
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run('Course Title')
            title_run.font.color.rgb = RGBColor(179, 0, 0)
            title_run.bold = True
            
            # Add course name
            course_para = course_title_cell.add_paragraph()
            course_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            course_run = course_para.add_run(metadata.course_title if metadata.course_title else '')
            course_run.font.color.rgb = RGBColor(179, 0, 0)
            course_run.bold = True
            
            # Add common text
            common_para = course_title_cell.add_paragraph()
            common_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # Add spacing

            # Part A
            partA_table = doc.add_table(rows=1, cols=5)
            partA_table.style = 'Table Grid'
            partA_table.allow_autofit = False  # Disable autofit
            
            # Set header row and column widths
            header_cells = partA_table.rows[0].cells
            headers = ['Q. No.', 'Questions', 'M', 'BT', 'CO']
            widths = [Inches(0.5), Inches(6.3), Inches(0.5), Inches(0.5), Inches(0.5)]
            
            # Set the column widths using cell properties
            for i, (text, width) in enumerate(zip(headers, widths)):
                cell = header_cells[i]
                # Set width in the cell properties
                tc = cell._tc
                tcW = tc.get_or_add_tcPr().get_or_add_tcW()
                tcW.type = 'dxa'
                tcW.w = width.twips
                
                QuestionPaperGenerator.format_table_cell(cell, text, WD_ALIGN_PARAGRAPH.CENTER)

            # Add Part A questions
            partA_questions = [q for q in selected_questions if q.part == 'A']
            partA_marks = QuestionPaperGenerator.add_questions_to_table(partA_questions, partA_table, widths, questions_data, doc)

            doc.add_paragraph()

            # Part B
            part_b_heading = doc.add_paragraph()
            part_b_run = part_b_heading.add_run('Part- B')
            part_b_run.bold = True
            part_b_run.font.size = Pt(12)
            part_b_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            partB_table = doc.add_table(rows=1, cols=5)
            partB_table.style = 'Table Grid'
            partB_table.allow_autofit = False
            
            header_cells = partB_table.rows[0].cells
            for i, (text, width) in enumerate(zip(headers, widths)):
                cell = header_cells[i]
                # Set width in the cell properties
                tc = cell._tc
                tcW = tc.get_or_add_tcPr().get_or_add_tcW()
                tcW.type = 'dxa'
                tcW.w = width.twips
                
                QuestionPaperGenerator.format_table_cell(cell, text, WD_ALIGN_PARAGRAPH.CENTER)

            # Add Part B questions
            partB_questions = [q for q in selected_questions if q.part == 'B']
            partB_marks = QuestionPaperGenerator.add_questions_to_table(partB_questions, partB_table, widths, questions_data, doc)

            # Add footer with total marks calculation
            doc.add_paragraph()
            footer_para = doc.add_paragraph("***")
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("BT-Blooms Taxonomy, CO-Course Outcomes")
            
            total_marks_para = doc.add_paragraph(f"Total Marks: {partA_marks + partB_marks}")
            total_marks_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            return doc

        except Exception as e:
            logging.error(f"Error generating paper: {str(e)}")
            raise Exception(f"Failed to generate paper: {str(e)}") 

    @staticmethod
    def create_answer_scheme(metadata, selected_questions):
        try:
            doc = Document()
            doc.add_heading('ANSWER SCHEME', 0)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.allow_autofit = False
            table.cell(0, 0).text = 'Q. No.'
            table.cell(0, 1).text = 'Answer'
            # Set Q. No. column width to 0.5 inches
            for row in table.rows:
                row.cells[0].width = Inches(0.5)

            for i, selection in enumerate(selected_questions, 1):
                q = selection.question
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                try:
                    qa = QuestionAnswer.objects.get(question=q)
                    row_cells[1].text = qa.answer
                except QuestionAnswer.DoesNotExist:
                    row_cells[1].text = 'N/A'

            # Explicitly set column widths for all rows after all rows are added
            for row in table.rows:
                row.cells[0].width = Inches(0.5)
                row.cells[1].width = Inches(6.0)

            return doc
        except Exception as e:
            logging.error(f"Error generating answer scheme: {str(e)}")
            raise Exception(f"Failed to generate answer scheme: {str(e)}")

    @staticmethod
    def add_questions_to_table(questions, table, widths, questions_data, doc):
        total_marks = 0
        for i, selection in enumerate(questions, 1):
            question_data = next(q for q in questions_data if q.q_id == selection.question.q_id)
            row_cells = table.add_row().cells
            
            # Set widths for new row
            for cell, width in zip(row_cells, widths):
                tc = cell._tc
                tcW = tc.get_or_add_tcPr().get_or_add_tcW()
                tcW.type = 'dxa'
                tcW.w = width.twips
            
            # Question number - keep it compact
            QuestionPaperGenerator.format_table_cell(row_cells[0], str(i), WD_ALIGN_PARAGRAPH.CENTER)
            row_cells[0].width = Inches(0.5)  # Fixed width for question number
            
            # Question content cell
            question_cell = row_cells[1]
            if not question_cell.text.strip():
                question_cell._element.clear_content()
            
            # Add question text
            para = question_cell.add_paragraph()
            para.add_run(question_data.text)
            
            # Add equations if present
            media = question_data.media.first()
            if media and media.equations:
                equations = media.equations
                for equation in equations:
                    eq_para = question_cell.add_paragraph()
                    QuestionPaperGenerator.add_equation_to_docx(eq_para, equation)
            
            # Add images if present
            if media and media.image_paths:
                for img_path in media.image_paths:
                    QuestionPaperGenerator.add_image_to_docx(doc, question_cell, img_path)
            
            # Add other fields
            QuestionPaperGenerator.format_table_cell(row_cells[2], str(question_data.marks), WD_ALIGN_PARAGRAPH.CENTER)
            QuestionPaperGenerator.format_table_cell(row_cells[3], question_data.bt, WD_ALIGN_PARAGRAPH.CENTER)
            QuestionPaperGenerator.format_table_cell(row_cells[4], question_data.co, WD_ALIGN_PARAGRAPH.CENTER)
            
            total_marks += question_data.marks
        
        return total_marks